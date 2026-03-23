from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

doc = Document()

# ── Style setup ──
style = doc.styles['Normal']
style.font.name = 'Malgun Gothic'
style.font.size = Pt(10)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Malgun Gothic'
    hs.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

# ── Helper functions ──
def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
    # Data
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table

def add_code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(8.5)
    return p

def add_quote(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(2)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.2 + level * 0.8)
    for r in p.runs:
        r.font.size = Pt(9.5)
    return p

# ══════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('JPX 그룹 자율규제 체계 분석 보고서')
run.bold = True
run.font.size = Pt(24)
run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('JPX-R과 TSE Listing Department의 조직·인력·역할 분담')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()
doc.add_paragraph()

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run('작성일: 2026-03-23\n').font.size = Pt(10)
meta.add_run('\n분석 대상 자료:\n').font.size = Pt(10)
r1 = meta.add_run('① E_20250130_1.pdf — Independent Directors\' Investigation Committee Report (2025.01.30)\n')
r1.font.size = Pt(9)
r2 = meta.add_run('② JPX-R_Annual_Report_2025_E.pdf — JPX-R Annual Report 2025 (FY2024)')
r2.font.size = Pt(9)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (manual)
# ══════════════════════════════════════════════════════════════
doc.add_heading('목차', level=1)
toc_items = [
    '1. JPX 그룹 전체 구조',
    '2. TSE Listing Department 상세 조직',
    '3. JPX-R 조직 및 인력 구성',
    '4. JPX-R 상장·공시 관련 업무 인원',
    '5. TSE Listing Department와 JPX-R의 역할 분담',
    '6. 전체 시장조치 — 판단주체·조치주체 종합표',
    '7. 감사의견 미표명(Disclaimer of Opinion) 처리 절차',
    '8. 내부자거래 사건과 정보관리 체계 개편',
    '9. 원본 자료 레퍼런스 가이드',
    '부록: 용어 대조표',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(1)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# SECTION 1
# ══════════════════════════════════════════════════════════════
doc.add_heading('1. JPX 그룹 전체 구조', level=1)
add_quote(doc, '출처: JPX-R_Annual_Report_2025_E.pdf pp.6–7 (Section I. Overview of Japan Exchange Regulation)')

add_table(doc,
    ['법인', '역할', '비고'],
    [
        ['Japan Exchange Group (JPX)', '지주회사', '상장회사 (8697, TSE Prime)'],
        ['Tokyo Stock Exchange (TSE)', '현물시장 운영', 'JPX 자회사'],
        ['Osaka Exchange (OSE)', '파생시장 운영', 'JPX 자회사'],
        ['Tokyo Commodity Exchange (TOCOM)', '파생시장 운영', 'JPX 자회사'],
        ['Japan Exchange Regulation (JPX-R)', '자율규제 (Self-Regulation)', '금상법상 독립 재단법인'],
        ['Japan Securities Clearing Corp. (JSCC)', '청산 (Clearing)', 'JPX 자회사'],
        ['JPX Market Innovation & Research', '시장관련 서비스', 'JPX 자회사'],
    ],
    col_widths=[5.5, 4, 5.5]
)

doc.add_heading('JPX-R의 법적 지위', level=2)
add_bullet(doc, '금융상품거래법(Financial Instruments and Exchange Act)에 의거하여 설립된 사단법인(membership association)')
add_bullet(doc, 'TSE 및 OSE로부터 자율규제 업무를 위탁(entrusted) 받아 수행')
add_bullet(doc, '이해충돌 방지를 위해 거래소(영리 운영)와 법적으로 독립된 운영이 의무화')
add_bullet(doc, '설립일: 2007년 10월 17일 (업무개시: 2007년 11월 1일)')
add_bullet(doc, '자본금: 30억엔 (JPY 3,000,000,000)')

doc.add_heading('그룹 구조도', level=2)
diagram1 = """┌─────────────────────────────────────────────────────────────┐
│               Japan Exchange Group, Inc. (JPX)              │
│           CEO: Yamaji Hiromi / COO: Iwanaga Moriyuki        │
│                    (상장: Code 8697, TSE Prime)              │
└──┬──────────┬──────────┬──────────┬──────────┬──────────────┘
   │          │          │          │          │
┌──▼──┐  ┌───▼───┐  ┌───▼───┐  ┌──▼───┐  ┌──▼──────────────┐
│ TSE │  │  OSE  │  │ TOCOM │  │ JSCC │  │    JPX-R        │
│현물 │  │ 파생  │  │ 파생  │  │ 청산 │  │  자율규제기관    │
│시장 │  │ 시장  │  │ 시장  │  │      │  │  (독립법인)      │
└─────┘  └───────┘  └───────┘  └──────┘  └─────────────────┘"""
add_code_block(doc, diagram1)

add_quote(doc, '추가 조사 포인트: JPX-R의 법적 독립성 요건 및 거버넌스에 대한 상세 내용은 JPX-R_Annual_Report_2025_E.pdf pp.6–8 참조')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# SECTION 2
# ══════════════════════════════════════════════════════════════
doc.add_heading('2. TSE Listing Department 상세 조직', level=1)
add_quote(doc, '출처: E_20250130_1.pdf pp.16–18 (Chapter 2, Section 3: Overview of the Listing Department)\n기준일: 2024년 10월 1일')

doc.add_heading('2.1 조직도', level=2)
diagram2 = """TSE Listing Department (총 82명)
│
├── Listing Department Director (1명) ── 부서 총괄
├── Section Director (1명) ── 부서 총괄 보좌
│
├─── [직할 3개 그룹] (26명)
│    ├── Planning Group (10명)
│    │     └─ 상장규칙 전반 기획
│    ├── Listed Company Support Group (4명)
│    │     └─ 상장기업 IR활동 지원
│    └── Administration Group (12명)
│          └─ 상장주식수 관리, 상장수수료 청구
│
└─── Corporate Disclosure Office (55명)
     │  Office장 (1명)
     │
     ├── Planning & Coordination Group (22명)
     │     ├─ 적시공시 규칙 기획
     │     ├─ Listed Company Services Group 지원
     │     └─ 감리종목 및 상폐종목 지정 공표
     │
     ├── Listed Company Services Group 1 — Prime Market (12명)
     │     └─ Prime Market 적시공시 접수·조언
     │
     ├── Listed Company Services Group 2 — Standard Market (13명)
     │     └─ Standard Market 적시공시 접수·조언
     │
     └── Listed Company Services Group 3 — Growth Market (7명)
           └─ Growth Market 적시공시 접수·조언"""
add_code_block(doc, diagram2)

doc.add_heading('2.2 인원 구성표', level=2)
add_table(doc,
    ['구분', '인원', '비고'],
    [
        ['Listing Department Director', '1명', '부서장'],
        ['Section Director', '1명', ''],
        ['Corporate Disclosure Office장', '1명', ''],
        ['Group Leaders', '9명', '직할 3그룹 + CDO 4그룹 + α'],
        ['직할 그룹 소계', '26명', ''],
        ['　Planning Group', '10명', '상장규칙 기획'],
        ['　Listed Company Support Group', '4명', 'IR 지원'],
        ['　Administration Group', '12명', '행정·수수료'],
        ['Corporate Disclosure Office 소계', '55명', ''],
        ['　Planning & Coordination Group', '22명', '공시규칙 기획, 감리/상폐종목 공표'],
        ['　LC Services Group 1 (Prime)', '12명', ''],
        ['　LC Services Group 2 (Standard)', '13명', ''],
        ['　LC Services Group 3 (Growth)', '7명', ''],
        ['합계', '82명', '정규직 77명 + 파견 5명'],
    ],
    col_widths=[6.5, 2, 6.5]
)

p = doc.add_paragraph()
r = p.add_run('Disclosure Supervisors (공시감독관): Listed Company Services Group 1~3 소속, 총 29명 (2024.10.1 기준). 이들이 TDnet을 통한 적시공시 접수·확인·공표 처리의 실무 주체.')
r.font.size = Pt(9.5)

add_quote(doc, '추가 조사 포인트: Listing Department 각 그룹 내부의 팀(team) 구성은 E_20250130_1.pdf pp.19–21에서 정보공유 범위 설명 시 일부 언급됨')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# SECTION 3
# ══════════════════════════════════════════════════════════════
doc.add_heading('3. JPX-R 조직 및 인력 구성', level=1)
add_quote(doc, '출처: JPX-R_Annual_Report_2025_E.pdf pp.4 (Introduction), pp.46 (Company Profile)')

doc.add_heading('3.1 기본 정보', level=2)
add_table(doc,
    ['항목', '내용'],
    [
        ['정식 명칭', 'Japan Exchange Regulation'],
        ['소재지', '2-1 Nihombashi Kabutocho, Chuo-ku, Tokyo 103-8229'],
        ['대표', 'President: Nakajima Junichi (나카지마 준이치)'],
        ['설립일', '2007년 10월 17일'],
        ['자본금', '30억엔 (JPY 3,000,000,000)'],
        ['최고 의사결정기구', 'Board of Governors (독립 governor 과반수)'],
    ],
    col_widths=[4, 11]
)

doc.add_heading('3.2 부서별 역할', level=2)
add_table(doc,
    ['부서', '주요 업무', '출처 페이지'],
    [
        ['Listing Examination Dept\n(상장심사부)', '신규상장 심사 (IPO, 시장구분 변경 등)', 'pp.11–14'],
        ['Listed Company Compliance Dept\n(상장회사 컴플라이언스부)', '적시공시 적정성 심사, 기업행동규범 준수 심사,\n상폐심사, 징계조치 결정', 'pp.15–22'],
        ['Trading Participants Exam &\nInspection Dept (거래참가자 검사부)', '증권사 등 거래참가자 검사·감독', 'pp.23–30'],
        ['Market Surveillance & Compliance\nDept (시장감시부)', '불공정거래 감시 (내부자거래, 시세조종)', 'pp.31–35'],
        ['General Administration Dept\n(총무부)', '일반관리, 보고서 편찬', 'p.51'],
    ],
    col_widths=[5, 7, 3]
)

doc.add_heading('3.3 JPX-R 조직 다이어그램', level=2)
diagram3 = """Japan Exchange Regulation (JPX-R)
│
├── Board of Governors (최고 의사결정기구, 독립 governor 과반수)
│
├── President: Nakajima Junichi
│
├── Listing Examination Dept ──────── 신규상장 심사
│     └─ IPO Liaison Meeting (TSE Listing Dept와 공동)
│
├── Listed Company Compliance Dept ── 상장유지 심사·징계
│     ├─ 적시공시 적정성 심사
│     ├─ 기업행동규범 준수 심사
│     ├─ 상폐기준 해당 심사
│     ├─ 특별주의시장종목(Special Alert) 지정 결정
│     ├─ 개선보고서 징구 / 위약금 부과 / 공표조치 결정
│     └─ Disciplinary Committee (자문기구)
│
├── Trading Participants Exam & Inspection Dept
│     ├─ 일반검사 / 추적검사 / 특별검사
│     └─ 거래자격 심사
│
├── Market Surveillance & Compliance Dept
│     ├─ 내부자거래 조사·심사
│     ├─ 시세조종 조사·심사
│     ├─ SESC 보고
│     └─ COMLEC (Compliance Learning Center) 운영
│
└── General Administration Dept"""
add_code_block(doc, diagram3)

p = doc.add_paragraph()
r = p.add_run('JPX 그룹 전체 임직원: 1,368명 (2024년 11월 서베이 기준, 출처: E_20250130_1.pdf p.11)')
r.font.size = Pt(9.5)
r.bold = True

add_quote(doc, '추가 조사 포인트: JPX-R 각 부서별 정확한 인원수는 두 자료에 명시되어 있지 않음. JPX IR자료 또는 유가증권보고서에서 세그먼트별 인원 확인 필요.')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# SECTION 4
# ══════════════════════════════════════════════════════════════
doc.add_heading('4. JPX-R 상장·공시 관련 업무 인원', level=1)
add_quote(doc, '출처: E_20250130_1.pdf pp.20–21, pp.31–32 (Chapter 2 Section 3(3)(c) 및 Chapter 4 Section 1(1)(b))')

doc.add_paragraph('TSE Listing Department의 Disclosure Supervisors로부터 미공개정보를 수령하는 정보수령자(information recipients) 등록 인원:')

doc.add_heading('4.1 미공개정보 수령 등록인원', level=2)
add_table(doc,
    ['JPX-R / TSE 부서', '변경 전 등록인원', '변경 후 (2025.1~)', '비고'],
    [
        ['TSE Equities Dept', '18명', '15명', '매매정지 등 시장운영 목적'],
        ['JPX-R Listed Company\nCompliance Dept', '25명', '7명', '변경 전: 부장 제외 전원\n→ 부서 총인원 26명+'],
        ['JPX-R Listing Examination Dept', '12명', '7명', '최근 상장기업 공시 확인 목적'],
        ['JPX-R Market Surveillance &\nCompliance Dept', '5명', '5명', '불공정거래 방지 목적, 변경 없음'],
    ],
    col_widths=[5, 3, 3, 5]
)

doc.add_heading('4.2 부서 규모 추정', level=2)
add_bullet(doc, 'Listed Company Compliance Dept: 변경 전 등록인원이 "부장 제외 전 부원"으로 25명 → 부서 총인원 최소 26명')
add_bullet(doc, 'Listing Examination Dept: 12명이 등록 수령자 → 부서 총인원은 12명 이상')
add_bullet(doc, 'Market Surveillance & Compliance Dept: 5명만 등록 → 전체 인원은 이보다 많을 것으로 추정')

add_quote(doc, '추가 조사 포인트: 정확한 부서별 인원은 원문(일본어판)이나 JPX-R 자체 공시자료에서 확인 필요.')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# SECTION 5
# ══════════════════════════════════════════════════════════════
doc.add_heading('5. TSE Listing Department와 JPX-R의 역할 분담', level=1)
add_quote(doc, '출처: JPX-R_Annual_Report_2025_E.pdf pp.6–8, pp.15–17; E_20250130_1.pdf pp.16–21')

doc.add_heading('핵심 원칙: 판단주체와 조치주체의 분리', level=2)
doc.add_paragraph('JPX 그룹의 자율규제 체계에서 가장 중요한 구조적 특징은 판단주체(Decision Maker)와 조치주체(Executor/Publisher)가 분리되어 있다는 점이다.')
add_quote(doc, '핵심 원문: "JPX-R independently conducts neutral inspections and makes decisions on approvals, disciplinary actions, and other measures under the name of the exchanges based on the results of these inspections." — JPX-R_Annual_Report_2025_E.pdf p.7')
add_bullet(doc, '판단주체: 심사·검토를 수행하고 조치의 필요성과 내용을 결정하는 기관')
add_bullet(doc, '조치주체: 판단주체의 결정에 따라 해당 조치를 공표·집행하는 기관')
doc.add_paragraph('대부분의 자율규제 조치에서 판단주체는 JPX-R, 조치주체는 TSE/OSE이다. 다만, 즉시적 시장운영 조치(매매정지 등)는 TSE가 판단과 조치를 모두 수행하며, 경고(Warning)·요청(Request) 등 일부 조치는 JPX-R이 판단과 조치를 모두 수행한다.')

doc.add_heading('5.1 역할 비교 — 판단주체 기준', level=2)
add_table(doc,
    ['구분', '판단주체', '조치주체', '비고'],
    [
        ['공시 접수·처리', 'TSE', 'TSE', 'Disclosure Supervisors가 TDnet 통해 수행'],
        ['공시 규칙 기획·제정', 'TSE', 'TSE', 'Planning Group / Planning & Coord. Group'],
        ['상장심사 (IPO)', 'JPX-R', 'TSE 명의', 'Listing Examination Dept 독립 심사'],
        ['감리종목 지정', 'JPX-R', 'TSE', 'JPX-R 심사 후 TSE가 공표'],
        ['상폐종목 지정', 'JPX-R', 'TSE', 'JPX-R 심사 후 TSE가 공표'],
        ['특별주의시장종목 지정', 'JPX-R', 'TSE 명의', 'Listed Co. Compliance Dept 결정'],
        ['개선보고서 징구', 'JPX-R', 'TSE 명의', 'Listed Co. Compliance Dept 결정'],
        ['위약금 부과', 'JPX-R', 'TSE 명의', 'Listed Co. Compliance Dept 결정'],
        ['상장폐지 (정량/정성)', 'JPX-R', 'TSE 명의', 'Listed Co. Compliance Dept 심사'],
        ['거래참가자 징계', 'JPX-R', 'TSE/OSE 명의', 'Disciplinary Committee 자문 후 결정'],
        ['경고 (상장기업 대상)', 'JPX-R', 'JPX-R', '판단·조치 모두 JPX-R'],
        ['경고 (거래참가자 대상)', 'JPX-R', 'JPX-R', '판단·조치 모두 JPX-R'],
        ['매매정지', 'TSE', 'TSE', 'Equities Dept 자체 판단·집행'],
        ['증거금거래 종목 선정', 'TSE', 'TSE', 'Equities Dept 자체 판단·집행'],
        ['불공정거래 감시', 'JPX-R', 'JPX-R→SESC', 'Market Surveillance Dept 전담'],
    ],
    col_widths=[4, 2.5, 3, 5.5]
)

doc.add_heading('5.2 정보 흐름도', level=2)
diagram_flow = """상장기업
  |
  | 적시공시 자료 제출 (TDnet)
  v
TSE Listing Dept -- Disclosure Supervisors (29명)
  |                         |
  | 공시 접수·확인·공표      | 미공개정보 공유 (사전등록 수령자)
  |                         |
  |              +----------+---------------------+
  |              |                                 |
  |        JPX-R Listed     JPX-R Listing    JPX-R Mkt Surv.
  |        Co. Compliance   Exam Dept        & Compliance
  |        (7명 등록)       (7명)            (5명)
  |              |
  |              | 심사 후 판단 결과 통보
  |<-------------+
  v
TSE 명의로 조치 공표·집행 (판단주체=JPX-R / 조치주체=TSE)"""
add_code_block(doc, diagram_flow)

doc.add_heading('5.3 독립성 확보 장치', level=2)
add_table(doc,
    ['장치', '내용', '출처'],
    [
        ['Board of Governors', '독립 governor 과반수로 구성', 'Annual Report p.7'],
        ['법적 분리', '금상법에 의거한 독립 법인 설립', 'Annual Report p.6'],
        ['독립적 의사결정', 'JPX-R이 독립적으로 심사·판단,\n거래소 명의로 집행', 'Annual Report p.7'],
        ['정기 평가', 'Board of Governors가\nJPX-R 실효성 정기 평가 수행', 'Annual Report p.8'],
    ],
    col_widths=[4, 7, 4]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# SECTION 6
# ══════════════════════════════════════════════════════════════
doc.add_heading('6. 전체 시장조치 — 판단주체·조치주체 종합표', level=1)
add_quote(doc, '출처: JPX-R_Annual_Report_2025_E.pdf 전체; E_20250130_1.pdf pp.16–21\n범례: 판단주체 = 심사·검토하여 조치를 결정하는 기관 / 조치주체 = 결정에 따라 공표·집행하는 기관')

doc.add_heading('6.1 상장기업에 대한 조치', level=2)
add_table(doc,
    ['조치', '판단주체', '조치주체', 'FY2024'],
    [
        ['감리종목 지정\n(Securities Under Supervision)', 'JPX-R\n(Listed Co. Compliance)', 'TSE\n(P&C Group 공표)', '-'],
        ['특별주의시장종목 지정\n(Security on Special Alert)', 'JPX-R\n(Listed Co. Compliance)', 'TSE 명의', '4건'],
        ['개선보고서 징구\n(Improvement Report)', 'JPX-R\n(Listed Co. Compliance)', 'TSE 명의', '7건'],
        ['상장계약위약금 부과\n(Violation Penalty)', 'JPX-R\n(Listed Co. Compliance)', 'TSE 명의', '5건'],
        ['공표조치\n(Public Announcement)', 'JPX-R\n(Listed Co. Compliance)', 'TSE 명의', '7건'],
        ['상폐 - 정량기준\n(상장유지기준 미달)', 'JPX-R\n(Listed Co. Compliance)', 'TSE 명의', '0건(심사)'],
        ['상폐 - 정량기준 기타\n(감사의견미표명 등)', 'JPX-R\n(Listed Co. Compliance)', 'TSE 명의', '106건(심사)'],
        ['상폐 - 정성기준\n(허위기재, 규칙위반)', 'JPX-R\n(Listed Co. Compliance)', 'TSE 명의', '1건(심사)'],
        ['상폐종목 지정\n(Security to Be Delisted)', 'JPX-R\n(Listed Co. Compliance)', 'TSE\n(P&C Group 공표)', '-'],
        ['실질적 존속회사 심사', 'JPX-R\n(Listed Co. Compliance)', 'TSE 명의', '88건(심사)'],
        ['경고 - 상장기업 대상\n(시장감시 기반)', 'JPX-R\n(Mkt Surveillance)', 'JPX-R\n(Governor/Dir/Mgr)', '2건'],
        ['내부체계 재검토 요청', 'JPX-R\n(Mkt Surveillance)', 'JPX-R', '5건'],
        ['특정거래 설명\n(상장기업 대상)', 'JPX-R\n(Mkt Surveillance)', 'JPX-R', '12건'],
    ],
    col_widths=[4, 3.5, 3.5, 2.5]
)

doc.add_heading('6.2 거래참가자(증권사 등)에 대한 조치', level=2)
add_table(doc,
    ['조치', '판단주체', '조치주체', 'FY2024'],
    [
        ['과징금 (Fine)', 'JPX-R\n+ Disciplinary Committee', 'TSE/OSE 명의', '3건'],
        ['견책 (Censure)', 'JPX-R\n+ Disciplinary Committee', 'TSE/OSE 명의', '3건'],
        ['매매정지/제한\n(Suspension)', 'JPX-R\n+ Disciplinary Committee', 'TSE/OSE 명의', '1건'],
        ['거래자격 취소\n(Cancellation)', 'JPX-R\n+ Disciplinary Committee', 'TSE/OSE 명의', '0건'],
        ['권고 (Recommendation)', 'JPX-R\n(Trading Participants Dept)', 'TSE 명의', '1건'],
        ['경고 - Governor 수준', 'JPX-R\n(Trading Participants Dept)', 'JPX-R (Governor)', '1건'],
        ['경고 - 부장 수준', 'JPX-R\n(Trading Participants Dept)', 'JPX-R (Dept Dir)', '4건'],
        ['경고 - 검사관 수준', 'JPX-R\n(Trading Participants Dept)', 'JPX-R (Inspector)', '8건'],
        ['요청 (Request)', 'JPX-R\n(Trading Participants Dept)', 'JPX-R', '7건'],
        ['특정거래 설명\n(거래참가자 대상)', 'JPX-R\n(Mkt Surveillance)', 'JPX-R', '299건'],
    ],
    col_widths=[4, 3.5, 3.5, 2.5]
)

doc.add_heading('6.3 즉시적 시장운영 조치 (TSE가 판단·조치 모두 수행)', level=2)
add_table(doc,
    ['조치', '판단주체', '조치주체', 'JPX-R 관여'],
    [
        ['매매정지 (Trading Halt)', 'TSE (Equities Dept)', 'TSE', 'X'],
        ['증거금거래 종목 선정', 'TSE (Equities Dept)', 'TSE', 'X'],
        ['실시간 매매 감시', 'TSE (Equities Dept)', 'TSE', 'X'],
        ['종목정보 업데이트', 'TSE (Equities Dept)', 'TSE', 'X'],
    ],
    col_widths=[5, 3.5, 2.5, 2.5]
)
add_quote(doc, '핵심 원문: "The TSE Equities Department is mostly in charge of market operations such as securities trading rules/trading halts, planning and drafting of margin trading rules, real-time surveillance of securities trading, and updating issue information." — E_20250130_1.pdf p.20')

doc.add_heading('6.4 심사·승인', level=2)
add_table(doc,
    ['조치', '판단주체', '조치주체', 'FY2024'],
    [
        ['신규상장 심사', 'JPX-R (Listing Exam Dept)', 'TSE 명의', '171건'],
        ['거래자격 심사', 'JPX-R', 'TSE/OSE 명의', '0건'],
        ['거래참가자 조직재편 심사', 'JPX-R', 'TSE/OSE 명의', '1건'],
    ],
    col_widths=[5, 3.5, 3, 2.5]
)

doc.add_heading('6.5 판단주체별 요약', level=2)
add_table(doc,
    ['판단주체', '해당 조치 유형'],
    [
        ['JPX-R\n(판단 + TSE 명의 집행)', '감리종목 지정, 상폐종목 지정, 특별주의시장종목 지정, 개선보고서 징구,\n위약금 부과, 공표조치, 상장폐지(정량/정성), 거래참가자 징계\n(과징금/견책/매매정지/자격취소), 신규상장 심사, 거래자격 심사'],
        ['JPX-R\n(판단 + JPX-R 자체 집행)', '경고(상장기업/거래참가자), 요청, 내부체계 재검토 요청,\n특정거래 설명, SESC 보고'],
        ['TSE\n(판단 + TSE 자체 집행)', '매매정지, 증거금거래 종목 선정, 실시간 매매 감시,\n종목정보 업데이트, 공시 접수·처리, 공시 규칙 기획'],
    ],
    col_widths=[4, 11]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# SECTION 7
# ══════════════════════════════════════════════════════════════
doc.add_heading('7. 감사의견 미표명(Disclaimer of Opinion) 처리 절차', level=1)
add_quote(doc, '출처: Securities Listing Regulations (2025.12.8) Rule 501, 503, 601, 608\nJPX-R_Annual_Report_2025_E.pdf pp.15-17; E_20250130_1.pdf pp.16-17\n규정 원문: https://www.jpx.co.jp/english/rules-participants/rules/regulations/index.html')

doc.add_heading('7.1 감리종목의 두 가지 서브카테고리', level=2)
doc.add_paragraph('TSE의 감리종목 지정에는 두 가지 서브카테고리가 있으며, 감사의견 미표명은 유예기간 없이 바로 지정되는 유형에 해당한다.')

add_table(doc,
    ['구분', 'Grace Period\n(猶予期間入り)', 'Under Confirmation\n(確認中)'],
    [
        ['트리거', '정량적 상장유지기준 미달\n(주주수, 유통주식 시가총액 등)', '감사의견 미표명/부적정,\n유가증권보고서 미제출 등'],
        ['Improvement Period', '있음 (원칙 1년)', '없음 - 바로 지정'],
        ['성격', '개선 기회 부여\n-> 기한 내 충족 시 해제', '상폐 해당 여부 확인 중\n-> 확인되면 정리종목->상폐'],
        ['근거 조항', 'Rule 501, Rule 601(1)', 'Rule 601(7)(8) 등'],
    ],
    col_widths=[3.5, 5.5, 5.5]
)

doc.add_heading('7.2 감사의견 미표명 시 경로 - TSE 트랙과 JPX-R 트랙', level=2)

p = doc.add_paragraph()
r = p.add_run('(A) TSE 행정 트랙 - 감리종목 지정 (Improvement Period 없음)')
r.bold = True
r.font.size = Pt(10)

diagram_tse = """감사의견 미표명 발생
       |
       v  TSE 행정 트랙 (유예기간 없음)
  監理銘柄(確認中) 지정      <- Improvement Period 없이 바로 지정
  Securities Under Supervision (Under Confirmation)
       |
       |  (투자자 고지 목적)
       v
  JPX-R 심사 결과에 따라:
  +-- 해소 (정정 감사의견 취득 등) -> 지정 해제
  +-- 상폐 확정 -> 整理銘柄(정리종목) -> 상장폐지"""
add_code_block(doc, diagram_tse)

add_quote(doc, '대비: 정량기준 미달(주주수 등)은 TSE 행정 트랙에서 監理銘柄(猶予期間入り)로 지정되어 1년의 improvement period가 부여된다.')

p = doc.add_paragraph()
r = p.add_run('(B) JPX-R 심사 트랙 - 두 가지 경로')
r.bold = True
r.font.size = Pt(10)

doc.add_paragraph('경로 1: 원칙적 경로 - 특별주의시장종목(Security on Special Alert) 지정 (Rule 503)')
diagram_jpxr1 = """JPX-R 심사: "내부관리체계 개선이 크게 필요한가?"
       |
       v Yes
  특별주의시장종목 지정 (Security on Special Alert)
  Rule 503 (1)(2)(b)
       |
       | 개선 기회 부여 (Written Confirmation 제출 등)
       |
   +-------+
   |       |
  개선완료 개선불가
   |       |
   v       v
  지정해제 Rule 601(9) -> 상장폐지"""
add_code_block(doc, diagram_jpxr1)

add_quote(doc, 'Rule 503 (1)(2)(b): "Where, in audit reports... a certified public accountant, etc. expresses an \'adverse opinion\' or \'disclaimer of opinion\'..." -> 내부관리체계 개선이 크게 필요하다고 판단 시 특별주의시장종목으로 지정')

doc.add_paragraph('경로 2: 극단적 경우 - 즉시 상장폐지 (Rule 601(8))')
diagram_jpxr2 = """JPX-R 심사: "즉시 상폐하지 않으면 시장 질서 유지가 명백히 곤란한가?"
       |
       v Yes (극히 예외적)
  Rule 601(8) -> 감리종목(確認中) -> 정리종목 -> 즉시 상장폐지"""
add_code_block(doc, diagram_jpxr2)

add_quote(doc, 'Rule 601 (1)(8): "Where such statement, etc. falls under Rule 503, Paragraph 1, Item (2) a. or b., and the Exchange deems that it clearly difficult to maintain order in the market if the listed company is not delisted immediately"')

doc.add_heading('7.3 TSE 트랙 vs JPX-R 트랙 비교', level=2)
add_table(doc,
    ['', 'TSE 행정 트랙', 'JPX-R 심사 트랙'],
    [
        ['역할', '감리종목(確認中) 지정·공표\n(투자자 고지)', '상폐기준 해당 여부 심사,\n조치 유형 결정'],
        ['Improvement\nPeriod', '없음 - 바로 지정', '경로에 따라 다름\n(Special Alert = 개선 기회 있음)'],
        ['판단주체', '- (규칙 기반 행정조치)', 'JPX-R (독립 심사)'],
        ['조치주체', 'TSE (P&C Group)', 'JPX-R(결정) -> TSE(공표)'],
        ['근거', 'Rule 608', 'Rule 503, Rule 601(8)(9)'],
    ],
    col_widths=[2.5, 5.5, 5.5]
)

doc.add_heading('7.4 정량기준 미달과의 비교', level=2)
add_table(doc,
    ['', '정량기준 미달 (주주수 등)', '감사의견 미표명'],
    [
        ['TSE 행정조치', '監理銘柄(猶予期間入り)', '監理銘柄(確認中)'],
        ['Improvement Period', '있음 (원칙 1년)', '없음 - 바로 지정'],
        ['성격', '개선 기회 부여', '상폐 해당 여부 확인'],
        ['이후 JPX-R 심사', '유예기간 만료 후 심사', '지정 즉시 심사 진행'],
        ['근거', 'Rule 501, Rule 601(1)', 'Rule 601(7)(8), Rule 503'],
    ],
    col_widths=[3.5, 5.5, 5.5]
)

doc.add_heading('7.5 FY2024 상폐심사 실적', level=2)
add_table(doc,
    ['심사 유형', '판단주체', 'FY2024', '전년 대비'],
    [
        ['Quantitative Criteria\n상장유지기준 미달', 'JPX-R', '0건', '+/-0'],
        ['Quantitative Criteria 기타\n(감사의견 미표명 등)', 'JPX-R', '106건', '+2'],
        ['Qualitative Criteria', 'JPX-R', '1건', '-2'],
        ['Substantial Surviving\nCompany 심사', 'JPX-R', '88건', '+29'],
    ],
    col_widths=[5, 3, 3, 3]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# SECTION 8
# ══════════════════════════════════════════════════════════════
doc.add_heading('8. 내부자거래 사건과 정보관리 체계 개편', level=1)
add_quote(doc, '출처: E_20250130_1.pdf 전체 (특히 Chapter 2–4)')

doc.add_heading('8.1 사건 개요', level=2)
add_bullet(doc, '2024년 9월, TSE Listing Department Corporate Disclosure Office 소속 직원이 SESC로부터 내부자거래 혐의로 수사')
add_bullet(doc, '공개매수(TOB) 관련 미공개정보 3건(Lawson, Riso Kyoiku, JASTEC)을 가족(생부)에게 전달')
add_bullet(doc, '2024년 12월 23일, SESC가 동경지검에 형사고발, 해당 직원 해고')

doc.add_heading('8.2 원인 분석', level=2)
doc.add_paragraph('조사위원회는 핵심 원인을 "need to share"가 "need to know"를 압도한 정보관리 체계로 지적:')

add_table(doc,
    ['문제점', '내용'],
    [
        ['부서 내 과도한 공유', '팀 내 전원, 그룹 내 전원, 부서 내 전원이\n미공개정보에 접근 가능'],
        ['공유폴더 무제한 접근', '각종 요약문서가 Listing Dept 전체 공유폴더에\n저장, 접근 제한 없음'],
        ['외부 부서 광범위 공유', 'Equities Dept 18명, Listed Co. Compliance 25명,\nListing Exam 12명, Mkt Surv. 5명'],
        ['근본 원인', '"직원이 내부자거래 규정을 위반할 가능성"을\n현실적 리스크로 인식하지 않았음'],
    ],
    col_widths=[4.5, 10.5]
)

doc.add_heading('8.3 재발방지 조치 (2025년 1월~)', level=2)
add_table(doc,
    ['조치', '변경 전', '변경 후'],
    [
        ['팀 내 공유 범위', '팀 전원 + 부서 전체 열람 가능', '담당 팀 + 매니저만'],
        ['Planning & Coordination Group', '전 부원에게 무조건 공유', '케이스 유형별 지정 인원 + 매니저만'],
        ['Equities Dept 수령자', '18명', '15명'],
        ['Listed Co. Compliance 수령자', '25명', '7명'],
        ['Listing Examination 수령자', '12명', '7명'],
        ['Market Surveillance 수령자', '5명', '5명 (변경 없음)'],
        ['공유범위 변경 절차', '규정 없음', 'Listing Dept Director 승인 필요'],
        ['3부서 공유 문서', '전원 접근 가능', '폐지'],
        ['내규 정비', '부서 자율', '새 정보관리 규칙 제정, 정기 점검'],
    ],
    col_widths=[5, 5, 5]
)

add_quote(doc, '추가 조사 포인트: 재발방지 조치의 구체적 실행 상황 및 위원회 평가는 E_20250130_1.pdf pp.31–41 (Chapter 4) 참조. 교육·연수 강화, 내규 개정, 커뮤니케이션 촉진 등 추가 조치도 기술되어 있음.')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# SECTION 9
# ══════════════════════════════════════════════════════════════
doc.add_heading('9. 원본 자료 레퍼런스 가이드', level=1)

doc.add_heading('9.1 E_20250130_1.pdf — 독립이사 조사위원회 보고서', level=2)
add_table(doc,
    ['섹션', '페이지', '핵심 내용', '추가 조사 활용'],
    [
        ['Chapter 1', 'pp.1–5', '조사위원회 설립 배경, 구성, 조사방법', '위원 구성, 법무자문 비용'],
        ['Chapter 2.3', 'pp.10–21', 'TSE Listing Dept 조직·인력·\n정보관리 체계 상세', '조직도, 인원수, 정보공유 범위의\n1차 출처'],
        ['Chapter 2.4', 'pp.16–17', '조사대상 직원의 경력·업무 상황', '채용·인사이동 프로세스'],
        ['Chapter 3', 'pp.18–24', '사건 원인 분석\n(정보관리, 가치관, 내규, 교육)', '구조적 문제점의 근본 원인 분석'],
        ['Chapter 4', 'pp.25–36', '재발방지 조치 및 위원회 평가', '정보공유 범위 축소 전후 비교,\n각 조치의 실효성 평가'],
        ['Chapter 5', 'p.37', '결론', '종합 평가'],
    ],
    col_widths=[3, 2.5, 5, 5]
)

doc.add_heading('9.2 JPX-R_Annual_Report_2025_E.pdf — JPX-R 연차보고서', level=2)
add_table(doc,
    ['섹션', '페이지', '핵심 내용', '추가 조사 활용'],
    [
        ['I. Overview', 'pp.3–8', 'JPX-R 조직구조, 법적 지위, 특징', '자율규제기관 설립 근거, 독립성 장치'],
        ['II. Changes', 'pp.7–9', '시장환경 변화와 대응', 'IPO Liaison Meeting,\n특별주의시장종목 제도'],
        ['III.1 Listing Exam', 'pp.11–14', '상장심사 업무, 심사건수, 결과', 'FY2024 신규상장 171건 심사 상세'],
        ['III.2 Listed Co.\nCompliance', 'pp.15–22', '상장유지 심사, 징계조치,\n특별주의시장종목 사례', 'FY2024 심사건수, 상폐·징계 실적,\n5개 Special Alert 사례'],
        ['III.3 Trading\nParticipants', 'pp.23–30', '거래참가자 검사·감독', '검사 24건, 징계 사례'],
        ['III.4 Market\nSurveillance', 'pp.31–35', '불공정거래 감시', '조사 2,958건, 심사 113건'],
        ['III.5 COMLEC', 'pp.36–40', '컴플라이언스 지원 활동', '세미나, e-러닝, 출판물'],
        ['III.6 Lists', 'pp.41–45', '신규상장·상폐·조치 목록', '개별 종목 리스트'],
        ['Company Profile', 'p.46', 'JPX-R 기본정보, 연혁', '설립일, 자본금, 연혁'],
    ],
    col_widths=[3, 2.5, 5, 5]
)

doc.add_heading('9.3 외부 참고 자료 (공식 웹사이트)', level=2)
add_table(doc,
    ['주제', 'URL'],
    [
        ['JPX-R 개요', 'https://www.jpx.co.jp/english/regulation/outline/about/index.html'],
        ['상장심사', 'https://www.jpx.co.jp/english/regulation/listing/eligibility/index.html'],
        ['상장유지 컴플라이언스', 'https://www.jpx.co.jp/english/regulation/listing/compliance/index.html'],
        ['거래참가자 검사', 'https://www.jpx.co.jp/english/regulation/maintaining/outline/index.html'],
        ['부정행위 방지 원칙', 'https://www.jpx.co.jp/english/regulation/listing/preventive-principles/index.html'],
        ['부정행위 대응 원칙', 'https://www.jpx.co.jp/english/regulation/listing/principle/index.html'],
        ['세미나·이벤트', 'https://www.jpx.co.jp/regulation/seminar/index.html (일본어)'],
        ['JPX 기업헌장', 'https://www.jpx.co.jp/english/corporate/governance/charter/index.html'],
    ],
    col_widths=[4.5, 11]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# APPENDIX
# ══════════════════════════════════════════════════════════════
doc.add_heading('부록: 용어 대조표', level=1)
add_table(doc,
    ['영문', '일본어 (추정)', '한국어'],
    [
        ['Securities Under Supervision', '監理銘柄', '감리종목'],
        ['Securities to Be Delisted', '整理銘柄', '정리종목 (상폐예정)'],
        ['Securities on Special Alert', '特設注意市場銘柄', '특별주의시장종목'],
        ['Timely Disclosure', '適時開示', '적시공시'],
        ['Disclosure Supervisor', '開示担当者', '공시감독관'],
        ['Listing Agreement Violation Penalty', '上場契約違約金', '상장계약위약금'],
        ['Improvement Report', '改善報告書', '개선보고서'],
        ['Public Announcement Measure', '公表措置', '공표조치'],
        ['Delisting Criteria', '上場廃止基準', '상장폐지기준'],
        ['Code of Corporate Conduct', '企業行動規範', '기업행동규범'],
        ['TDnet', '適時開示情報伝達システム', '적시공시정보전달시스템'],
        ['SESC', '証券取引等監視委員会', '증권거래등감시위원회'],
        ['COMLEC', 'コンプライアンス学習センター', '컴플라이언스학습센터'],
    ],
    col_widths=[5.5, 5, 4.5]
)

# ── Footer note ──
doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run('본 문서는 E_20250130_1.pdf와 JPX-R_Annual_Report_2025_E.pdf의 영문 번역본을 기반으로 작성되었습니다. 원문(일본어)과 차이가 있을 경우 일본어 원문이 우선합니다.')
r.italic = True
r.font.size = Pt(8)
r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

# ── Save ──
output_path = r'c:\Users\Peter\Desktop\temp\JPX_JPX-R_TSE_조직및역할분석.docx'
doc.save(output_path)
print(f'Saved: {output_path}')
